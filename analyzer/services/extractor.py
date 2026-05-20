import pandas as pd
import pdfplumber
import docx
import re
import io

def compute_data_profile(df):
    """Computes statistical profile of a DataFrame to enable metric validation."""
    import numpy as np
    numeric_df = df.select_dtypes(include='number')
    n = len(df)
    profile = {
        'n': n,
        'has_negative': False,
        'has_zero': False,
        'unique_ratio': 0.0,
        'has_numeric': len(numeric_df.columns) > 0,
    }
    if len(numeric_df.columns) > 0:
        flat = numeric_df.values.flatten()
        flat = flat[~np.isnan(flat.astype(float))]
        if len(flat) > 0:
            profile['has_negative'] = bool(np.any(flat < 0))
            profile['has_zero'] = bool(np.any(flat == 0))
            profile['unique_ratio'] = round(float(len(np.unique(flat)) / len(flat)), 4)
    return profile

def extract_docx_dataframes(uploaded_file):
    """Parses Word (.docx) tables into Pandas DataFrames, mimicking Excel sheets."""
    import numpy as np
    import pandas as pd
    result_dfs = {}
    try:
        uploaded_file.seek(0)
        file_bytes = uploaded_file.read()
        doc = docx.Document(io.BytesIO(file_bytes))
        
        # 1. Parse tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if not table_data:
                continue
            
            # Pad rows to the max number of columns
            max_cols = max(len(r) for r in table_data)
            for idx, r in enumerate(table_data):
                if len(r) < max_cols:
                    table_data[idx] = r + [''] * (max_cols - len(r))
            
            if len(table_data) > 1:
                headers = table_data[0]
                seen = {}
                clean_headers = []
                for h_idx, h in enumerate(headers):
                    h_str = h if h else f"Columna_{h_idx+1}"
                    if h_str in seen:
                        seen[h_str] += 1
                        clean_headers.append(f"{h_str}_{seen[h_str]}")
                    else:
                        seen[h_str] = 1
                        clean_headers.append(h_str)
                df = pd.DataFrame(table_data[1:], columns=clean_headers)
            elif len(table_data) == 1:
                headers = [f"Columna_{col_idx+1}" for col_idx in range(max_cols)]
                df = pd.DataFrame(table_data, columns=headers)
            else:
                continue
            
            # Convert columns to numeric if possible
            for col in df.columns:
                try:
                    df[col] = pd.to_numeric(df[col])
                except (ValueError, TypeError):
                    pass
            
            sheet_name = f"Tabla {i+1}"
            result_dfs[sheet_name] = df
            
        # 2. Fallback: parse text paragraphs if no tables are found
        if not result_dfs:
            lines = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    lines.append(t)
            if lines:
                df = pd.DataFrame({'Contenido': lines})
                result_dfs["Texto"] = df
                
    except Exception as e:
        print("Error en extract_docx_dataframes:", e)
        
    return result_dfs

def extract_text_metrics(text_content):
    """
    Analiza texto plano, guarda el texto real de cada párrafo y devuelve un DataFrame 
    con métricas lingüísticas por párrafo para mapeo de variables en el frontend.
    """
    from collections import Counter
    import re
    import pandas as pd

    STOP_WORDS = {
        'el', 'la', 'los', 'las', 'un', 'una', 'unos', 'unas', 'de', 'del', 'al',
        'en', 'y', 'o', 'a', 'que', 'es', 'se', 'no', 'con', 'por', 'para',
        'su', 'sus', 'lo', 'como', 'más', 'pero', 'le', 'ya', 'este', 'esta',
        'han', 'ha', 'he', 'si', 'me', 'mi', 'tu', 'te', 'nos', 'son', 'fue',
        'ser', 'también', 'todo', 'todos', 'cada', 'sobre', 'entre', 'cuando',
        'muy', 'sin', 'hasta', 'desde', 'durante', 'después', 'antes', 'aunque',
        'mientras', 'mediante', 'hacia', 'través', 'embargo', 'bien', 'tan',
        'vez', 'veces', 'porque', 'sino', 'donde', 'quien', 'cual', 'esto',
        'ello', 'sido', 'estar', 'tiene', 'tienen', 'puede', 'pueden', 'parte',
        'hace', 'hacer', 'así', 'mismo', 'solo', 'ellos', 'ellas', 'nosotros',
        'ante', 'bajo', 'cabe', 'contra', 'tras', 'versus', 'via', 'pro',
    }

    paragraphs = [p.strip() for p in text_content.split('\n') if p.strip()]
    if not paragraphs:
        raise ValueError("El archivo de texto está vacío o no tiene contenido legible.")

    all_text = ' '.join(paragraphs)
    all_words = re.findall(r'\b[a-záéíóúüñ]{4,}\b', all_text.lower())
    filtered = [w for w in all_words if w not in STOP_WORDS]
    top_keywords = [word for word, _ in Counter(filtered).most_common(8)]

    rows = []
    for i, para in enumerate(paragraphs):
        words = re.findall(r'\b\w+\b', para.lower())
        sentences = [s.strip() for s in re.split(r'[.!?]+', para) if s.strip()]
        unique = set(words)
        para_kw_counts = Counter(re.findall(r'\b[a-záéíóúüñ]{4,}\b', para.lower()))

        row = {
            'Párrafo':         i + 1,
            'Contenido':       para,  # <--- Enviamos el texto real del párrafo
            'Total_Palabras':  len(words),
            'Total_Oraciones': len(sentences),
            'Palabras_Únicas': len(unique),
            'Densidad_Léxica': round(len(unique) / len(words), 3) if words else 0,
        }
        for kw in top_keywords:
            row[f'Frec_{kw}'] = para_kw_counts.get(kw, 0)
        rows.append(row)

    return pd.DataFrame(rows)


def get_file_info(uploaded_file):
    filename = uploaded_file.name
    filename_lower = filename.lower()
    result_list = []
    
    try:
        if filename_lower.endswith(('.xls', '.xlsx')):
            import numpy as np
            # Leemos todas las hojas de forma eficiente
            dfs = pd.read_excel(uploaded_file, sheet_name=None, nrows=100)
            for sheet_name, df in dfs.items():
                df = df.replace({pd.NA: None, np.nan: None})
                profile = compute_data_profile(df)
                result_list.append({
                    'name': f"{filename} - {sheet_name}",
                    'columns': [str(col) for col in df.columns],
                    'preview': df.to_dict('records'),
                    'profile': profile
                })
        elif filename_lower.endswith(('.doc', '.docx')):
            import numpy as np
            dfs = extract_docx_dataframes(uploaded_file)
            for sheet_name, df in dfs.items():
                df = df.replace({pd.NA: None, np.nan: None})
                profile = compute_data_profile(df)
                result_list.append({
                    'name': f"{filename} - {sheet_name}",
                    'columns': [str(col) for col in df.columns],
                    'preview': df.to_dict('records')[:100],
                    'profile': profile
                })
        elif filename_lower.endswith('.txt'):
            import numpy as np
            uploaded_file.seek(0)
            text_content = uploaded_file.read().decode('utf-8', errors='ignore')
            df = extract_text_metrics(text_content)
            df_clean = df.replace({pd.NA: None, np.nan: None})
            profile = compute_data_profile(df_clean)
            result_list.append({
                'name': filename,
                'columns': [str(col) for col in df.columns],
                'preview': df_clean.to_dict('records'),
                'profile': profile,
                'txt_info': True  # Marca para el frontend
            })
        else:
            result_list.append({
                'name': filename,
                'columns': ['Datos Genéricos del Documento'],
                'preview': [{'Dato': 'Vista previa no disponible para este formato'}],
                'profile': {'n': 0, 'has_negative': False, 'has_zero': False, 'unique_ratio': 0.0, 'has_numeric': False}
            })
    except Exception as e:
        print("Error en get_file_info:", e)
        
    return result_list

def extract_numbers(uploaded_file, target_column=None):
    filename = uploaded_file.name.lower()
    numbers = []

    try:
        if filename.endswith(('.xls', '.xlsx')):
            try:
                dfs = pd.read_excel(uploaded_file, sheet_name=None)
                
                for sheet_name, df in dfs.items():
                    # Extraer de todas las hojas que contengan la columna objetivo, o aplanar toda la hoja
                    if target_column and target_column in df.columns:
                        source_values = df[target_column].dropna().values
                    elif not target_column:
                        source_values = df.values.flatten()
                    else:
                        continue # La columna no está en esta hoja

                    for value in source_values:
                        try:
                            num = float(value)
                            if not pd.isna(num):
                                numbers.append(num)
                        except (ValueError, TypeError):
                            pass
            except Exception as e:
                raise ValueError(f"Error al leer Excel multicapa. Verifica que el archivo no esté corrupto. Detalle: {str(e)}")

        elif filename.endswith('.pdf'):
            try:
                file_bytes = uploaded_file.read()
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            matches = re.findall(r'-?\d+\.?\d*', text)
                            numbers.extend([float(x) for x in matches])
            except Exception as e:
                raise ValueError(f"Error al procesar el PDF. Asegúrate de que contiene texto leíble. Detalle: {str(e)}")
        elif filename.endswith('.txt'):
            try:
                uploaded_file.seek(0)
                text_content = uploaded_file.read().decode('utf-8', errors='ignore')
                df = extract_text_metrics(text_content)
                if target_column and target_column in df.columns:
                    col_values = df[target_column].dropna().tolist()
                    numbers.extend([float(v) for v in col_values])
                else:
                    # Sin columna objetivo: usar Total_Palabras por párrafo
                    numbers.extend(df['Total_Palabras'].tolist())
            except Exception as e:
                raise ValueError(f"Error al analizar el texto para extraer métricas numéricas. Detalle: {str(e)}")
        elif filename.endswith(('.doc', '.docx')):
            try:
                dfs = extract_docx_dataframes(uploaded_file)
                found_in_table = False
                if target_column:
                    for sheet_name, df in dfs.items():
                        if target_column in df.columns:
                            source_values = df[target_column].dropna().values
                            for value in source_values:
                                try:
                                    num = float(value)
                                    if not pd.isna(num):
                                        numbers.append(num)
                                        found_in_table = True
                                except (ValueError, TypeError):
                                    pass
                
                if not found_in_table:
                    uploaded_file.seek(0)
                    doc = docx.Document(uploaded_file)
                    for paragraph in doc.paragraphs:
                        matches = re.findall(r'-?\d+\.?\d*', paragraph.text)
                        numbers.extend([float(x) for x in matches])
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                matches = re.findall(r'-?\d+\.?\d*', cell.text)
                                numbers.extend([float(x) for x in matches])
            except Exception as e:
                raise ValueError(f"Error al procesar el documento Word. Verifica su formato y contenido. Detalle: {str(e)}")
        
        else:
            raise ValueError("Formato de archivo no compatible. Solo se permiten archivos Excel (.xlsx, .xls), Word (.docx, .doc), PDF o texto plano (.txt).")

        # Limpiar y retornar si existe informacion
        if not numbers:
            raise ValueError("El archivo subido no contiene cifras o datos numéricos válidos accesibles. Por favor verifica su contenido.")
        
        return numbers

    except ValueError as ve:
        raise ve
    except Exception as e:
        raise ValueError(f"Ocurrió un error inesperado durante la extracción de datos: {str(e)}")


def extract_categorical_column(uploaded_file, target_column):
    """
    Extrae la secuencia de estados para las Cadenas de Markov.
    Si es un archivo de texto (.txt), devuelve la secuencia lineal de palabras clave del artículo.
    """
    import pandas as pd
    filename = uploaded_file.name.lower()
    values = []

    try:
        if filename.endswith(('.xls', '.xlsx')):
            dfs = pd.read_excel(uploaded_file, sheet_name=None)
            for sheet_name, df in dfs.items():
                if target_column in df.columns:
                    values.extend(df[target_column].dropna().astype(str).tolist())
            return values

        elif filename.endswith(('.doc', '.docx')):
            dfs = extract_docx_dataframes(uploaded_file)
            for sheet_name, df in dfs.items():
                if target_column in df.columns:
                    values.extend(df[target_column].dropna().astype(str).tolist())
            return values

        elif filename.endswith('.txt'):
            uploaded_file.seek(0)
            text_content = uploaded_file.read().decode('utf-8', errors='ignore')
            
            import re
            STOP_WORDS = {'el', 'la', 'los', 'las', 'de', 'del', 'en', 'que', 'y', 'a', 'un', 'una', 'es', 'se', 'para', 'por', 'su'}
            # Extrae todas las palabras significativas correlativas del ensayo
            all_words = re.findall(r'\b[a-záéíóúüñ]{4,}\b', text_content.lower())
            values = [w for w in all_words if w not in STOP_WORDS]
            
            if not values:
                raise ValueError("El archivo .txt no produjo palabras clave válidas para la secuencia.")
            return values
            
        else:
            raise ValueError("Formato no compatible.")
    except Exception as e:
        raise ValueError(f"Error al extraer datos categóricos: {str(e)}")

